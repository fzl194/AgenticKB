"""Structural DOCX parser using python-docx.

Strategy:
1. Open the DOCX file via python-docx Document.
2. Walk body elements in order (paragraphs + tables).
3. Detect headings by style name (Heading 1..6).
4. Extract embedded drawings (a:blip / rId) as block_type=image.
5. Build a SectionNode tree using the same stack-based builder as pdf_parser.
"""
from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from knowledge_mining.mining.contracts.models import ContentBlock, SectionNode

logger = logging.getLogger(__name__)

_HEADING_STYLE_RE = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)
_LIST_STYLE_RE = re.compile(r"^List", re.IGNORECASE)
_BULLET_PREFIX_RE = re.compile(r"^[•\-\*\u2022\u25CF\u25CB]\s+")
_CAPTION_RE = re.compile(
    r"^(?:图|表|Figure|Table|Abbildung)\s*[\d.]+",
    re.IGNORECASE,
)


def parse_docx_to_section_tree(
    file_path: str,
    doc_title: str | None = None,
    *,
    image_dir: str | None = None,
) -> SectionNode:
    """Parse a DOCX file into a SectionNode tree."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed; cannot parse DOCX files")
        return SectionNode(title=doc_title, level=0)

    doc = Document(file_path)
    blocks = _extract_blocks(doc, image_dir=image_dir)
    if not blocks:
        return SectionNode(title=doc_title, level=0)
    return _build_section_tree(blocks, doc_title)


def count_embedded_images(file_path: str) -> int:
    """Count image-like related parts in a DOCX package (for .doc conversion checks).

    Prefers package relationships (robust across OOXML variants) and falls back
    to counting ``a:blip`` embeds in the body.
    """
    try:
        from docx import Document
    except ImportError:
        return 0
    try:
        doc = Document(file_path)
    except Exception:
        return 0
    n = 0
    try:
        for rel in doc.part.rels.values():
            reltype = getattr(rel, "reltype", "") or ""
            if "image" in reltype:
                n += 1
    except Exception:
        n = 0
    if n:
        return n
    return _count_blips_in_element(doc.element.body)


def _count_blips_in_element(element: Any) -> int:
    from docx.oxml.ns import qn

    try:
        return len(element.findall(".//" + qn("a:blip")))
    except Exception:
        return 0


def _extract_blocks(doc: Any, *, image_dir: str | None) -> list[ContentBlock]:
    """Walk document body elements in order, yielding ContentBlocks."""
    blocks: list[ContentBlock] = []
    image_idx = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            paragraph = None
            for p in doc.paragraphs:
                if p._element is element:
                    paragraph = p
                    break
            if paragraph is None:
                continue
            para_blocks, image_idx = _blocks_from_paragraph(
                paragraph, doc=doc, image_dir=image_dir, image_idx=image_idx,
            )
            blocks.extend(para_blocks)

        elif tag == "tbl":
            table = None
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
            if table is not None:
                table_block, img_blocks, image_idx = _extract_table_with_images(
                    table, doc=doc, image_dir=image_dir, image_idx=image_idx,
                )
                blocks.append(table_block)
                blocks.extend(img_blocks)

    return blocks


def _blocks_from_paragraph(
    para: Any,
    *,
    doc: Any,
    image_dir: str | None,
    image_idx: int,
) -> tuple[list[ContentBlock], int]:
    """Emit text block (if any) plus any embedded images in document order."""
    from docx.oxml.ns import qn

    out: list[ContentBlock] = []
    text = para.text.strip()
    images: list[ContentBlock] = []

    for blip in para._element.findall(".//" + qn("a:blip")):
        block, image_idx = _image_block_from_blip(
            blip, doc=doc, image_dir=image_dir, image_idx=image_idx,
            native_caption=text if _CAPTION_RE.match(text or "") else None,
        )
        if block is not None:
            images.append(block)

    classified = _classify_paragraph_text(text, para) if text else None
    # Pure image paragraph: keep images, drop empty text.
    if classified is not None:
        out.append(classified)
    out.extend(images)
    return out, image_idx


def _classify_paragraph_text(text: str, para: Any) -> ContentBlock | None:
    if not text:
        return None
    style_name = (para.style.name or "") if para.style else ""

    m = _HEADING_STYLE_RE.match(style_name)
    if m:
        level = int(m.group(1))
        level = max(1, min(6, level))
        return ContentBlock(block_type="heading", text=text, level=level)

    is_list = _LIST_STYLE_RE.match(style_name) is not None if style_name else False
    if is_list or _BULLET_PREFIX_RE.match(text):
        return ContentBlock(block_type="list", text=text)

    return ContentBlock(block_type="paragraph", text=text)


def _extract_table_with_images(
    table: Any,
    *,
    doc: Any,
    image_dir: str | None,
    image_idx: int,
) -> tuple[ContentBlock, list[ContentBlock], int]:
    """Extract table text; append any cell-embedded images after the table block."""
    from docx.oxml.ns import qn

    rows_text: list[str] = []
    col_count = 0
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        col_count = max(col_count, len(cells))
        rows_text.append("| " + " | ".join(cells) + " |")

    if len(rows_text) >= 2:
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows_text.insert(1, separator)

    table_text = "\n".join(rows_text)
    table_block = ContentBlock(
        block_type="table",
        text=table_text,
        structure={
            "kind": "table",
            "row_count": len(table.rows),
            "col_count": col_count,
        },
    )

    images: list[ContentBlock] = []
    for blip in table._element.findall(".//" + qn("a:blip")):
        block, image_idx = _image_block_from_blip(
            blip, doc=doc, image_dir=image_dir, image_idx=image_idx,
        )
        if block is not None:
            images.append(block)
    return table_block, images, image_idx


def _image_block_from_blip(
    blip: Any,
    *,
    doc: Any,
    image_dir: str | None,
    image_idx: int,
    native_caption: str | None = None,
) -> tuple[ContentBlock | None, int]:
    from docx.oxml.ns import qn
    from knowledge_mining.mining.infra.image_assets import materialize_image_bytes

    embed = blip.get(qn("r:embed"))
    if not embed:
        return None, image_idx
    try:
        part = doc.part.related_parts[embed]
        data = part.blob
    except Exception as exc:
        logger.debug("DOCX image part %s unreadable: %s", embed, exc)
        return None, image_idx
    if not data:
        return None, image_idx

    content_type = getattr(part, "content_type", "") or ""
    ext = _ext_for_content_type(content_type)
    structure: dict[str, Any] = {
        "kind": "docx_image",
        "rId": embed,
        "content_type": content_type,
    }
    if native_caption:
        structure["native_caption"] = native_caption

    if image_dir:
        meta = materialize_image_bytes(
            data, image_dir, stem=f"docx_{image_idx:02d}", ext=ext,
        )
        structure["image_path"] = meta["image_path"]
        structure["image_sha256"] = meta["image_sha256"]
    else:
        import hashlib

        structure["image_sha256"] = hashlib.sha256(data).hexdigest()
        structure["caption_source"] = "missing_image_dir"

    image_idx += 1
    return ContentBlock(block_type="image", text="", structure=structure), image_idx


def _ext_for_content_type(content_type: str) -> str:
    mapping = {
        "image/png": ".png",
        "image/jpeg": ".jpg",
        "image/jpg": ".jpg",
        "image/gif": ".gif",
        "image/bmp": ".bmp",
        "image/webp": ".webp",
        "image/x-emf": ".emf",
        "image/x-wmf": ".wmf",
        "image/tiff": ".tiff",
    }
    return mapping.get((content_type or "").lower(), ".bin")


def _build_section_tree(
    blocks: list[ContentBlock], doc_title: str | None,
) -> SectionNode:
    """Stack-based builder: pop ancestors with level >= current heading level."""
    root = _mutable_node(title=doc_title, level=0)
    stack: list[dict] = [root]

    for block in blocks:
        if block.block_type == "heading" and block.level:
            level = block.level
            while len(stack) > 1 and stack[-1]["level"] >= level:
                stack.pop()
            new_node = _mutable_node(title=block.text, level=level)
            stack[-1]["children"].append(new_node)
            stack.append(new_node)
        else:
            stack[-1]["blocks"].append(block)

    return _freeze(root)


def _mutable_node(title: str | None, level: int) -> dict:
    return {"title": title, "level": level, "children": [], "blocks": []}


def _freeze(node: dict) -> SectionNode:
    return SectionNode(
        title=node["title"],
        level=node["level"],
        blocks=tuple(node["blocks"]),
        children=tuple(_freeze(c) for c in node["children"]),
    )
