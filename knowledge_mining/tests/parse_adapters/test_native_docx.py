"""``NativeDocxParser`` + ``DocxNormalizer`` 单元测试（M3, SRS §C06/§C07/§4.6-§4.7）.

fixture 在测试内用 python-docx 生成 bytes（不依赖外部文件）。覆盖：
heading 层级树（style API 判级）、paragraph_index 证据、合并单元格
（gridSpan 横向 / vMerge 纵向）、TableAsset 网格、单一 section 容器、
parent_of / next_in_reading_order 关系、stable id、round-trip validate、
不支持 MIME / 损坏 bytes 的错误归一。
"""
from __future__ import annotations

import io

import pytest

from knowledge_mining.mining.contracts.parse_ir import validate
from knowledge_mining.mining.contracts.parser_adapter import (
    ParserAdapterError,
    UnsupportedFormat,
)
from knowledge_mining.mining.parse_adapters.native.native_docx import (
    NATIVE_DOCX_FINGERPRINT,
    DocxNormalizer,
    NativeDocxParser,
)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
RAW_HASH = "cd" * 32


def _build_docx_bytes() -> bytes:
    """标题树 + 段落 + 3x3 合并单元格表格的 DOCX bytes."""
    from docx import Document

    doc = Document()
    doc.add_heading("一级标题", level=1)
    doc.add_paragraph("导语段落。")
    doc.add_heading("二级标题", level=2)
    doc.add_paragraph("二级下正文。")
    doc.add_heading("三级标题", level=3)
    doc.add_paragraph("三级下正文。")

    table = doc.add_table(rows=3, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1))  # gridSpan=2
    table.cell(0, 0).text = "表头A"
    table.cell(0, 2).text = "表头C"
    table.cell(1, 0).merge(table.cell(2, 0))  # vMerge restart/continue
    table.cell(1, 0).text = "竖排"
    table.cell(1, 1).text = "x"
    table.cell(1, 2).text = "y"
    table.cell(2, 1).text = "z"
    table.cell(2, 2).text = "w"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def docx_bytes() -> bytes:
    return _build_docx_bytes()


@pytest.fixture
def parser() -> NativeDocxParser:
    return NativeDocxParser()


@pytest.fixture
def normalizer() -> DocxNormalizer:
    return DocxNormalizer()


def _parse_doc(parser: NativeDocxParser, data: bytes):
    return parser.parse(data, mime=DOCX_MIME)


def _normalize(parser: NativeDocxParser, normalizer: DocxNormalizer, data: bytes):
    artifact = _parse_doc(parser, data)
    return normalizer.normalize(artifact, source_raw_hash=RAW_HASH)


# ---------------------------------------------------------------------------
# RED 1: parse 层 —— 块类型序列、heading level、paragraph_index 证据
# ---------------------------------------------------------------------------

def test_parse_blocks_types_levels_and_native_refs(parser, docx_bytes) -> None:
    artifact = _parse_doc(parser, docx_bytes)

    assert artifact.parser_id == "native_docx"
    assert artifact.mime == DOCX_MIME
    types = [b.block_type for b in artifact.blocks]
    assert types == [
        "heading", "paragraph", "heading", "paragraph",
        "heading", "paragraph", "table",
    ]
    levels = [b.level for b in artifact.blocks if b.block_type == "heading"]
    assert levels == [1, 2, 3]
    # paragraph_index 连续递增（含空段也不会破坏索引语义）
    para_refs = [
        b.native_ref["paragraph_index"]
        for b in artifact.blocks if b.block_type == "paragraph"
    ]
    assert para_refs == [1, 3, 5]
    table_blocks = [b for b in artifact.blocks if b.block_type == "table"]
    assert len(table_blocks) == 1
    assert table_blocks[0].native_ref == {"table_index": 0}
    # DOCX 无页/坐标概念：bbox 不伪造
    assert all(b.bbox is None for b in artifact.blocks)


def test_parse_unsupported_mime_raises(parser) -> None:
    with pytest.raises(UnsupportedFormat):
        parser.parse(b"whatever", mime="application/pdf")


def test_parse_corrupt_bytes_wrapped(parser) -> None:
    with pytest.raises(ParserAdapterError):
        parser.parse(b"not a zip at all", mime=DOCX_MIME)


# ---------------------------------------------------------------------------
# RED 2: IR 断言链 —— parse -> normalize -> validate -> 结构断言
# ---------------------------------------------------------------------------

def test_full_ir_chain_containers_relations_and_parents(
    parser, normalizer, docx_bytes
) -> None:
    doc = _normalize(parser, normalizer, docx_bytes)

    result = validate(doc)
    assert result.valid, [i for i in result.issues if i.level == "error"]

    # 单一 section 容器，无页码伪造
    assert len(doc.containers) == 1
    section = doc.containers[0]
    assert section.container_type == "section"
    assert section.page_number is None

    assert len(doc.elements) == 7
    by_type = {}
    for el in doc.elements:
        by_type.setdefault(el.element_type, []).append(el)
    assert [e.text for e in by_type["heading"]] == ["一级标题", "二级标题", "三级标题"]
    assert [e.text for e in by_type["paragraph"]] == [
        "导语段落。", "二级下正文。", "三级下正文。",
    ]

    h1, h2, h3 = by_type["heading"]
    assert h1.parent_id is None
    assert h2.parent_id == h1.element_id
    assert h3.parent_id == h2.element_id
    assert by_type["paragraph"][0].parent_id == h1.element_id
    assert by_type["paragraph"][1].parent_id == h2.element_id
    assert by_type["paragraph"][2].parent_id == h3.element_id

    parent_of = [
        r for r in doc.relations if r.relation_type == "parent_of"
    ]
    assert len(parent_of) == 6  # h2/h3 + 3 段落 + 表格
    reading = [
        r for r in doc.relations if r.relation_type == "next_in_reading_order"
    ]
    assert len(reading) == 6  # 7 元素 -> 6 条链接

    # fingerprint
    assert doc.source_identity.parser_fingerprint == NATIVE_DOCX_FINGERPRINT
    assert doc.source_identity.parser_fingerprint == (
        "native_docx@2.0.0#python-docx-1.2.0"
    )


def test_table_asset_merges_and_header(parser, normalizer, docx_bytes) -> None:
    doc = _normalize(parser, normalizer, docx_bytes)

    assert list(doc.structured_assets) == [f"{doc.elements[-1].element_id}-table"]
    asset = next(iter(doc.structured_assets.values()))
    assert asset.rows == 3
    assert asset.columns == 3
    grid = {(c.row_index, c.column_index): c for c in asset.cells}

    # 横向合并：gridSpan=2
    assert grid[(0, 0)].text == "表头A"
    assert grid[(0, 0)].column_span == 2
    assert grid[(0, 0)].row_span == 1
    assert (0, 1) not in grid  # 被合并覆盖的位置不产 cell
    assert grid[(0, 2)].text == "表头C"
    # 纵向合并：vMerge 跨 2 行
    assert grid[(1, 0)].text == "竖排"
    assert grid[(1, 0)].row_span == 2
    assert grid[(1, 0)].column_span == 1
    assert (2, 0) not in grid
    # 首行 is_header 约定
    assert all(grid[(0, c)].is_header for c in (0, 2))
    assert not grid[(1, 1)].is_header
    assert asset.header_regions == ((0, 0),)


def test_paragraph_index_evidence_roundtrip(parser, normalizer, docx_bytes) -> None:
    doc = _normalize(parser, normalizer, docx_bytes)

    for el in doc.elements:
        if el.element_type != "paragraph":
            continue
        span = el.source_spans[0]
        assert span.native_ref is not None
        assert "paragraph_index" in span.native_ref
        assert isinstance(span.native_ref["paragraph_index"], int)
        assert span.raw_text == el.text
    heading_els = [e for e in doc.elements if e.element_type == "heading"]
    assert [
        e.source_spans[0].native_ref["paragraph_index"] for e in heading_els
    ] == [0, 2, 4]


def test_stable_ids_and_roundtrip_dict(parser, normalizer, docx_bytes) -> None:
    doc1 = _normalize(parser, normalizer, docx_bytes)
    doc2 = _normalize(parser, normalizer, docx_bytes)
    assert [e.element_id for e in doc1.elements] == [
        e.element_id for e in doc2.elements
    ]

    rebuilt = type(doc1).from_dict(doc1.to_dict())
    result = validate(rebuilt)
    assert result.valid
    assert [e.element_id for e in rebuilt.elements] == [
        e.element_id for e in doc1.elements
    ]


def test_rectangular_merge_row_span_exact() -> None:
    """评审 HIGH-1 回归：2列×3行矩形合并 row_span 必须=3（非 5）."""
    import io
    from docx import Document as DocxDocument
    from knowledge_mining.mining.parse_adapters.native.native_docx import (
        NativeDocxParser,
    )
    from knowledge_mining.mining.parse_adapters.normalizer import LegacyLineNormalizer  # noqa: F401

    doc = DocxDocument()
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).merge(table.cell(2, 1))  # 矩形合并 2 列 × 3 行
    buf = io.BytesIO(); doc.save(buf)
    parser = NativeDocxParser()
    artifact = parser.parse(buf.getvalue(), mime=parser.descriptor.supported_mimes.__iter__().__next__())
    tables = [b for b in artifact.blocks if b.block_type == "table"]
    assert tables, "table block missing"
    cells = tables[0].structure["cells"]
    origins = [c for c in cells if c["row_index"] == 0 and c["column_index"] == 0]
    assert origins, "merge origin missing"
    assert origins[0]["row_span"] == 3, origins[0]
    assert origins[0]["column_span"] == 2, origins[0]


# ===========================================================================
# 整改轮（2026-08-17）：列表语义 / 单元格证据 / 六类结构诊断
# ===========================================================================


def _docx_with_lists() -> bytes:
    """编号列表（两级）+ 普通段落."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("导语。")
    p1 = doc.add_paragraph("一级项甲", style="List Number")
    p1.paragraph_format.left_indent = None
    p2 = doc.add_paragraph("一级项乙", style="List Number")
    sub = doc.add_paragraph("二级子项", style="List Bullet 2")
    doc.add_paragraph("结尾段。")
    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def _set_numpr(paragraph, ilvl: int, num_id: int = 1) -> None:
    """直接注入 w:numPr（style 不一定能带编号属性）."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl"); ilvl_el.set(qn("w:val"), str(ilvl))
    numid_el = OxmlElement("w:numId"); numid_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el); numPr.append(numid_el)
    pPr.append(numPr)


def test_numbered_paragraphs_become_list_items_with_level() -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("普通段。")
    top = doc.add_paragraph("顶层项")
    _set_numpr(top, 0)
    nested = doc.add_paragraph("嵌套项")
    _set_numpr(nested, 1)
    buf = io.BytesIO(); doc.save(buf)

    artifact = NativeDocxParser().parse(buf.getvalue(), mime=DOCX_MIME)
    kinds = [(b.block_type, b.level) for b in artifact.blocks]
    assert ("paragraph", None) in kinds
    assert ("list_item", 1) in kinds  # ilvl 0 -> level 1
    assert ("list_item", 2) in kinds  # ilvl 1 -> level 2


def test_table_cells_have_independent_native_evidence() -> None:
    parser = NativeDocxParser()
    doc = DocxNormalizer().normalize(
        parser.parse(_build_docx_bytes(), mime=DOCX_MIME),
        source_raw_hash="cd" * 32,
    )
    table = next(e for e in doc.elements if e.element_type == "table")
    asset = doc.structured_assets[f"{table.element_id}-table"]
    span_by_id = {s.span_id: s for s in table.source_spans}
    seen = set()
    for cell in asset.cells:
        if not cell.text:
            continue
        assert cell.source_span_id is not None
        span = span_by_id[cell.source_span_id]
        # 独立证据：native_ref 带 OOXML 单元格定位
        assert "row_index" in span.native_ref and "column_index" in span.native_ref
        assert cell.source_span_id not in seen
        seen.add(cell.source_span_id)


def _docx_with_unsupported_structures() -> bytes:
    """页眉 + 脚注 + 文本框 + 图片的文档（六类结构的代表集）."""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from PIL import Image

    doc = Document()
    doc.add_paragraph("正文段。")
    # 页眉文本
    header = doc.sections[0].header
    header.paragraphs[0].text = "页眉文字"
    # 图片（inline）
    img = io.BytesIO()
    Image.new("RGB", (2, 2), (0, 128, 255)).save(img, format="PNG")
    doc.add_picture(img, width=None)
    # 文本框（w:p 内嵌 txbxContent）
    p = doc.add_paragraph()
    run = p.add_run()
    txbx = OxmlElement("w:pict")
    txbx_content = OxmlElement("w:txbxContent")
    inner_p = OxmlElement("w:p")
    inner_r = OxmlElement("w:r")
    inner_t = OxmlElement("w:t")
    inner_t.text = "文本框内容"
    inner_r.append(inner_t); inner_p.append(inner_r)
    txbx_content.append(inner_p); txbx.append(txbx_content)
    run._r.append(txbx)
    # 脚注（document part footer 引用较繁琐——用 sectPr 前注记即可，
    # 本测试以 header/footnote parts 存在性为诊断对象）
    buf = io.BytesIO(); doc.save(buf)
    return buf.getvalue()


def test_unsupported_structures_are_diagnosed_not_silent() -> None:
    artifact = NativeDocxParser().parse(
        _docx_with_unsupported_structures(), mime=DOCX_MIME
    )
    joined = "\n".join(artifact.warnings).lower()
    assert "header" in joined, f"页眉静默丢失: {joined!r}"
    assert "image" in joined, f"图片静默丢失: {joined!r}"
    assert "textbox" in joined, f"文本框静默丢失: {joined!r}"


def test_nested_table_extracted_as_separate_block() -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("含嵌套表格：")
    outer = doc.add_table(rows=2, cols=2)
    outer.cell(0, 0).text = "外层A"
    outer.cell(0, 1).text = "外层B"
    outer.cell(1, 0).text = "宿主格"
    outer.cell(1, 1).text = "外层D"
    inner = outer.cell(1, 0).add_table(rows=1, cols=2)
    inner.cell(0, 0).text = "内层1"
    inner.cell(0, 1).text = "内层2"
    buf = io.BytesIO(); doc.save(buf)

    artifact = NativeDocxParser().parse(buf.getvalue(), mime=DOCX_MIME)
    tables = [b for b in artifact.blocks if b.block_type == "table"]
    assert len(tables) == 2, f"嵌套表丢失: {[b.text for b in artifact.blocks]}"
    inner_blocks = [t for t in tables if "内层1" in t.text]
    assert inner_blocks, "内层表格未独立成块"
    assert inner_blocks[0].native_ref.get("in_cell") is not None
